import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuration : ajustez ces chemins si besoin
INPUT_FILE = r'C:\Users\GLT\eiffage.com\OneDrive - eiffageenergie.be\Documents\Guillaume\Programmation\Auto. Doc. PBI\fichier_converti.json'
OUTPUT_FILE = r'C:\Users\GLT\eiffage.com\OneDrive - eiffageenergie.be\Documents\Guillaume\Programmation\Auto. Doc. PBI\documentation.docx'
EXCLUDE_KEYWORDS = ['LocalDateTable', 'DateTableTemplate']

def load_report(path):
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)

def extract_metadata(report):
    tables = []
    calc_tables = []
    for tbl in report.get('model', {}).get('tables', []):
        name = tbl.get('name', '')
        # Exclusion
        if any(key in name for key in EXCLUDE_KEYWORDS):
            continue
        # Table calculée
        if tbl.get('type') == 'calculatedTable':
            expr = tbl.get('expression', '')
            if isinstance(expr, list):
                expr = '\n'.join(expr)
            calc_tables.append({'name': name, 'expression': expr})
            continue
        # Table standard
        info = {
            'name': name,
            'measures': [],
            'calculated_columns': [],
            'partitions': [],
            'hierarchies': []
        }
        # Mesures
        for m in tbl.get('measures', []):
            expr = m.get('expression', '')
            if isinstance(expr, list):
                expr = '\n'.join(expr)
            info['measures'].append({'name': m.get('name'), 'expression': expr})
        # Colonnes calculées
        for col in tbl.get('columns', []):
            if col.get('type') in ['calculated', 'calculatedTableColumn']:
                expr = col.get('expression', '')
                if isinstance(expr, list):
                    expr = '\n'.join(expr)
                info['calculated_columns'].append({'name': col.get('name'), 'expression': expr})
        # Partitions (incluant type et expression)
        for part in tbl.get('partitions', []):
            src = part.get('source', {})
            src_type = src.get('type', '')
            src_expr = src.get('expression', '') if src_type == 'calculated' else ''
            info['partitions'].append({
                'name': part.get('name'),
                'mode': part.get('mode'),
                'source_type': src_type,
                'source_expression': src_expr
            })
        # Hiérarchies
        for hier in tbl.get('hierarchies', []):
            levels = [lvl.get('name') for lvl in hier.get('levels', [])]
            info['hierarchies'].append({'name': hier.get('name'), 'levels': levels})
        tables.append(info)
    return tables, calc_tables

def style_table(table, headers, rows):
    """
    Applique un style soigné à une table Word :
    - En-têtes en Calibri 11 pt blanc sur fond rouge, bordures fines noires.
    - Alternance de lignes blanc/gris clair.
    - Colonne d'expressions en Consolas 9 pt monospace, indentation, interligne resserré.
    """
    # ——— 1. Style de l’en-tête ———
    hdr_cells = table.rows[0].cells
    for col_idx, header_text in enumerate(headers):
        cell = hdr_cells[col_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        # Police et couleur du texte de l'en-tête
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        # Fond rouge
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'FF0000')
        tcPr.append(shd)
        # Bordures fines noires
        borders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            bd = OxmlElement(f'w:{side}')
            bd.set(qn('w:val'), 'single')
            bd.set(qn('w:sz'), '4')       # épaisseur 0.5 pt
            bd.set(qn('w:space'), '0')
            bd.set(qn('w:color'), '000000')
            borders.append(bd)
        tcPr.append(borders)

    # ——— 2. Style des lignes de données ———
    for row_idx, row_values in enumerate(rows):
        cells = table.add_row().cells
        # Couleur de fond alternée
        bg_color = 'FFFFFF' if (row_idx % 2 == 0) else 'F2F2F2'

        for col_idx, cell_value in enumerate(row_values):
            cell = cells[col_idx]
            cell.text = ''
            lines = str(cell_value).split('\n')
            for line_idx, text_line in enumerate(lines):
                if line_idx == 0:
                    p = cell.paragraphs[0]
                else:
                    p = cell.add_paragraph()
                run = p.add_run(text_line)

                # ——— 2.a Colonne d'expression (indice 1) : Consolas 9 pt, indentation ———
                if col_idx == 1:
                    run.font.name = 'Consolas'
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Consolas')
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p.paragraph_format.left_indent = Pt(6)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.1
                else:
                    # Colonne non-expression : Calibri 10 pt, texte noir
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.15

            # ——— 2.b Fond et bordures pour chaque cellule ———
            tcPr = cell._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), bg_color)
            tcPr.append(shd)

            borders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                bd = OxmlElement(f'w:{side}')
                bd.set(qn('w:val'), 'single')
                bd.set(qn('w:sz'), '4')
                bd.set(qn('w:space'), '0')
                bd.set(qn('w:color'), '000000')
                borders.append(bd)
            tcPr.append(borders)

def generate_word(tables, calc_tables, output_path):
    doc = Document()
    doc.add_heading('Documentation automatique du modèle Power BI', level=0)

    # --- Tables calculées ---
    if calc_tables:
        doc.add_heading('Tables calculées', level=1)
        t = doc.add_table(rows=1, cols=2)
        headers = ['Nom', 'Expression']
        rows = [(c['name'], c['expression']) for c in calc_tables]
        style_table(t, headers, rows)
        doc.add_paragraph()

    # --- Tables usuelles ---
    for tbl in tables:
        doc.add_heading(f"Table: {tbl['name']}", level=1)

        # Mesures
        if tbl['measures']:
            doc.add_paragraph('Mesures :', style='Intense Quote')
            t = doc.add_table(rows=1, cols=2)
            headers = ['Nom', 'Expression']
            rows = [(m['name'], m['expression']) for m in tbl['measures']]
            style_table(t, headers, rows)
            doc.add_paragraph()

        # Colonnes calculées
        if tbl['calculated_columns']:
            doc.add_paragraph('Colonnes calculées :', style='Intense Quote')
            t = doc.add_table(rows=1, cols=2)
            headers = ['Nom', 'Expression']
            rows = [(c['name'], c['expression']) for c in tbl['calculated_columns']]
            style_table(t, headers, rows)
            doc.add_paragraph()

        # Partitions
        if tbl['partitions']:
            doc.add_paragraph('Partitions :', style='Intense Quote')
            t = doc.add_table(rows=1, cols=4)
            headers = ['Nom', 'Mode', 'Type source', 'Expression source']
            rows = [
                (p['name'], p['mode'], p['source_type'], p['source_expression'])
                for p in tbl['partitions']
            ]
            style_table(t, headers, rows)
            doc.add_paragraph()

        # Hiérarchies
        if tbl['hierarchies']:
            doc.add_paragraph('Hiérarchies :', style='Intense Quote')
            for hier in tbl['hierarchies']:
                doc.add_paragraph(f"{hier['name']}: {', '.join(hier['levels'])}")
            doc.add_paragraph()

    doc.save(output_path)
    print(f"Documentation Word générée : {output_path}")

if __name__ == '__main__':
    report = load_report(INPUT_FILE)
    tables, calc_tables = extract_metadata(report)
    generate_word(tables, calc_tables, OUTPUT_FILE)
