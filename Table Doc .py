import json
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === 🔧 Répertoire de base dynamique ===
base_dir = Path(__file__).resolve().parent
layout_file = base_dir / "Layout.json"
output_file = base_dir / "documentation.docx"


# === 📥 Chargement JSON Layout ===
with open(layout_file, encoding='utf-8') as f:
    layout = json.load(f)

# === 🧠 Extraction des informations de page ===
pages = []
for section in layout.get("sections", []):
    name = section.get("displayName", "Sans nom")
    page_id = section.get("name", "Inconnu")
    config_str = section.get("config")
    visibility = "Non défini (None)"
    if config_str:
        try:
            config_data = json.loads(config_str)
            vis_val = config_data.get("visibility")
            if vis_val == 0:
                visibility = "Visible"
            elif vis_val == 1:
                visibility = "Masquée"
            else:
                visibility = f"Non défini ({vis_val})"
        except json.JSONDecodeError:
            visibility = "Erreur JSON"
    pages.append({"Nom de la page": name, "ID de la page": page_id, "Visibilité": visibility})

df_pages = pd.DataFrame(pages)

# === 🖨 Affichage console ===
print("\n=== Tableau des pages Power BI ===\n")
print(df_pages.to_string(index=False))

# === 🎨 Fonction de style de tableau Word identique ===
def style_table(table, headers, rows):
    # En-têtes stylés
    hdr_cells = table.rows[0].cells
    for col_idx, header_text in enumerate(headers):
        cell = hdr_cells[col_idx]
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

        # Fond rouge
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'FF0000')
        tcPr.append(shd)

        # Bordures
        borders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            bd = OxmlElement(f'w:{side}')
            bd.set(qn('w:val'), 'single')
            bd.set(qn('w:sz'), '4')
            bd.set(qn('w:space'), '0')
            bd.set(qn('w:color'), '000000')
            borders.append(bd)
        tcPr.append(borders)

    # Lignes de données
    for row_idx, row_values in enumerate(rows):
        cells = table.add_row().cells
        bg_color = 'FFFFFF' if (row_idx % 2 == 0) else 'F2F2F2'
        for col_idx, val in enumerate(row_values):
            cell = cells[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15

            # Fond + bordures
            tcPr = cell._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), bg_color)
            tcPr.append(shd)

            borders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                bd = OxmlElement(f'w:{side}')
                bd.set(qn('w:val'), 'single')
                bd.set(qn('w:sz'), '4')
                bd.set(qn('w:space'), '0')
                bd.set(qn('w:color'), '000000')
                borders.append(bd)
            tcPr.append(borders)

# === 📝 Insertion du tableau dans le document Word ===
doc = Document(output_file)
doc.add_page_break()
doc.add_heading("Visibilité des pages", level=1)

headers = list(df_pages.columns)
rows = df_pages.values.tolist()
table = doc.add_table(rows=1, cols=len(headers))
style_table(table, headers, rows)

doc.save(output_file)
print(f"\n✅ Tableau ajouté dans le fichier Word : {output_file}")
